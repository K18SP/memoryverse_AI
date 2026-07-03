"""
DOCX Parser
───────────
Extracts paragraphs and table cells from .docx files.
Tables are flattened to tab-separated rows so the chunker
can still split them sensibly.
"""

import io
from docx import Document


async def extract_docx(raw_bytes: bytes, filename: str) -> dict:
    doc    = Document(io.BytesIO(raw_bytes))
    parts  = []

    # ── Paragraphs ─────────────────────────────────────────────────────────
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    # ── Tables ─────────────────────────────────────────────────────────────
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return {
        "text"      : "\n".join(parts),
        "extra_meta": {
            "paragraph_count": len(doc.paragraphs),
            "table_count"    : len(doc.tables),
        },
    }
